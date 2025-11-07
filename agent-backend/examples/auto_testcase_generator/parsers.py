"""解析器工具"""
from pathlib import Path
from typing import List, Tuple
from functools import lru_cache
import requests

from .models import ApiEndpoint, BusinessScenario


class SwaggerParser:
    """Swagger解析器"""
    
    @staticmethod
    @lru_cache(maxsize=10)
    def parse(url: str, max_endpoints: int = 20) -> Tuple[List[ApiEndpoint], List[BusinessScenario]]:
        """解析Swagger URL（带缓存）"""
        swagger_json = requests.get(url, timeout=30).json()
        endpoints = SwaggerParser._parse_endpoints(swagger_json, max_endpoints)
        scenarios = SwaggerParser._identify_scenarios(endpoints)
        return endpoints, scenarios
    
    @staticmethod
    def _parse_endpoints(swagger_json: dict, max_endpoints: int) -> List[ApiEndpoint]:
        """解析接口"""
        endpoints = []
        for path, methods in swagger_json.get('paths', {}).items():
            for method, spec in methods.items():
                if method.upper() not in ['GET', 'POST', 'PUT', 'DELETE', 'PATCH']:
                    continue
                endpoints.append(ApiEndpoint(
                    path=path,
                    method=method.upper(),
                    operation_id=spec.get('operationId', f"{method}_{path}"),
                    summary=spec.get('summary', ''),
                    description=spec.get('description', ''),
                    parameters=spec.get('parameters', []),
                    request_body=spec.get('requestBody'),
                    responses=spec.get('responses', {}),
                    tags=spec.get('tags', [])
                ))
                if len(endpoints) >= max_endpoints:
                    return endpoints
        return endpoints
    
    @staticmethod
    def _identify_scenarios(endpoints: List[ApiEndpoint]) -> List[BusinessScenario]:
        """识别业务场景"""
        tag_groups = {}
        for ep in endpoints:
            for tag in ep.tags:
                tag_groups.setdefault(tag, []).append(ep)
        
        scenarios = []
        for tag, eps in tag_groups.items():
            crud_ops = {'create': [], 'read': [], 'update': [], 'delete': []}
            for ep in eps:
                op = ep.operation_id.lower()
                if ep.method == 'POST' or any(k in op for k in ['create', 'add']):
                    crud_ops['create'].append(ep)
                elif ep.method == 'GET' or any(k in op for k in ['get', 'find', 'list']):
                    crud_ops['read'].append(ep)
                elif ep.method in ['PUT', 'PATCH'] or 'update' in op:
                    crud_ops['update'].append(ep)
                elif ep.method == 'DELETE' or 'delete' in op:
                    crud_ops['delete'].append(ep)
            
            scenario_eps = [ops[0] for ops in crud_ops.values() if ops]
            if len(scenario_eps) >= 2:
                scenarios.append(BusinessScenario(
                    name=f"{tag}_完整业务流程",
                    description=f"{tag}模块的端到端测试（CRUD）",
                    endpoints=scenario_eps,
                    execution_order=[ep.operation_id for ep in scenario_eps]
                ))
        return scenarios


class DocumentParser:
    """文档解析器"""
    
    @staticmethod
    def parse(file_path: str) -> str:
        """解析文档"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        parsers = {
            '.txt': DocumentParser._parse_txt,
            '.docx': DocumentParser._parse_word,
            '.pdf': DocumentParser._parse_pdf,
        }
        
        parser = parsers.get(path.suffix.lower())
        if not parser:
            raise ValueError(f"不支持的格式: {path.suffix}")
        
        return parser(path)
    
    @staticmethod
    def _parse_txt(path: Path) -> str:
        """解析TXT"""
        for encoding in ['utf-8', 'gbk', 'gb2312']:
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        raise UnicodeDecodeError(f"无法解码: {path}")
    
    @staticmethod
    def _parse_word(path: Path) -> str:
        """解析Word"""
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = [' | '.join(cell.text.strip() for cell in row.cells)
                 for table in doc.tables for row in table.rows]
        return '\n'.join(paragraphs + tables)
    
    @staticmethod
    def _parse_pdf(path: Path) -> str:
        """解析PDF"""
        from pypdf import PdfReader
        return '\n'.join(page.extract_text() for page in PdfReader(path).pages
                        if page.extract_text().strip())

