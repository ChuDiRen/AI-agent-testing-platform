"""
文档解析器基类和实现
"""
import os
import hashlib
from abc import ABC, abstractmethod
from typing import List, Dict, Any
from pathlib import Path
import asyncio
from dataclasses import dataclass

from core.logger import setup_logger

logger = setup_logger(__name__)


@dataclass
class ParsedDocument:
    """解析后的文档"""
    text: str
    metadata: Dict[str, Any]
    pages: List[str] = None
    tables: List[Dict[str, Any]] = None
    images: List[str] = None


class DocumentParser(ABC):
    """文档解析器抽象基类"""

    @abstractmethod
    def parse(self, file_path: str) -> ParsedDocument:
        """解析文档

        Args:
            file_path: 文件路径

        Returns:
            解析后的文档
        """
        pass

    @abstractmethod
    def supports(self, file_type: str) -> bool:
        """是否支持该文件类型

        Args:
            file_type: 文件类型（pdf, docx, txt等）

        Returns:
            是否支持
        """
        pass


class PDFParser(DocumentParser):
    """PDF文档解析器"""

    def __init__(self):
        try:
            import pypdf
            self.pypdf = pypdf
            logger.info("PDF解析器初始化成功")
        except ImportError:
            logger.warning("pypdf未安装，PDF解析功能不可用")
            self.pypdf = None

    def supports(self, file_type: str) -> bool:
        return file_type.lower() == "pdf"

    def parse(self, file_path: str) -> ParsedDocument:
        """解析PDF文档"""
        if not self.pypdf:
            raise ImportError("请安装 pypdf: pip install pypdf")

        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        text_parts = []
        pages = []
        metadata = {
            "file_type": "pdf",
            "file_name": file_path_obj.name,
            "file_size": file_path_obj.stat().st_size,
        }

        try:
            with open(file_path, "rb") as f:
                pdf_reader = self.pypdf.PdfReader(f)

                # 提取元数据
                if pdf_reader.metadata:
                    metadata.update({
                        "title": pdf_reader.metadata.get("/Title", ""),
                        "author": pdf_reader.metadata.get("/Author", ""),
                        "creator": pdf_reader.metadata.get("/Creator", ""),
                        "producer": pdf_reader.metadata.get("/Producer", ""),
                        "page_count": len(pdf_reader.pages)
                    })
                else:
                    metadata["page_count"] = len(pdf_reader.pages)

                # 提取每一页的文本
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        pages.append(page_text)
                        text_parts.append(page_text)

                # 合并所有文本
                full_text = "\n\n".join(text_parts)

                # 计算文本哈希
                text_hash = hashlib.md5(full_text.encode()).hexdigest()
                metadata["text_hash"] = text_hash
                metadata["char_count"] = len(full_text)
                metadata["word_count"] = len(full_text.split())

                logger.info(f"PDF解析成功: {file_path_obj.name} (页数: {len(pages)}, 字符数: {len(full_text)})")

                return ParsedDocument(
                    text=full_text,
                    metadata=metadata,
                    pages=pages,
                    tables=[],
                    images=[]
                )

        except Exception as e:
            logger.error(f"PDF解析失败: {file_path}, 错误: {str(e)}")
            raise


class DocxParser(DocumentParser):
    """Word文档解析器"""

    def __init__(self):
        try:
            import docx
            self.docx = docx
            logger.info("Word解析器初始化成功")
        except ImportError:
            logger.warning("python-docx未安装，Word解析功能不可用")
            self.docx = None

    def supports(self, file_type: str) -> bool:
        return file_type.lower() in ["docx", "doc"]

    def parse(self, file_path: str) -> ParsedDocument:
        """解析Word文档"""
        if not self.docx:
            raise ImportError("请安装 python-docx: pip install python-docx")

        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        try:
            doc = self.docx.Document(file_path)

            # 提取段落文本
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            full_text = "\n".join(paragraphs)

            # 提取表格
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables.append({
                        "rows": len(table_data),
                        "cols": len(table_data[0]) if table_data else 0,
                        "data": table_data
                    })

            # 提取元数据
            metadata = {
                "file_type": "docx",
                "file_name": file_path_obj.name,
                "file_size": file_path_obj.stat().st_size,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }

            # 计算文本哈希
            text_hash = hashlib.md5(full_text.encode()).hexdigest()
            metadata["text_hash"] = text_hash
            metadata["char_count"] = len(full_text)
            metadata["word_count"] = len(full_text.split())

            logger.info(f"Word文档解析成功: {file_path_obj.name} (段落: {len(doc.paragraphs)}, 表格: {len(doc.tables)})")

            return ParsedDocument(
                text=full_text,
                metadata=metadata,
                pages=[],
                tables=tables,
                images=[]
            )

        except Exception as e:
            logger.error(f"Word文档解析失败: {file_path}, 错误: {str(e)}")
            raise


class TxtParser(DocumentParser):
    """纯文本解析器"""

    def supports(self, file_type: str) -> bool:
        return file_type.lower() == "txt"

    def parse(self, file_path: str) -> ParsedDocument:
        """解析纯文本文件"""
        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'big5', 'latin1']
            text = None
            used_encoding = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                        used_encoding = encoding
                        break
                except UnicodeDecodeError:
                    continue

            if text is None:
                raise ValueError("无法解码文件，尝试了多种编码都失败")

            # 提取元数据
            metadata = {
                "file_type": "txt",
                "file_name": file_path_obj.name,
                "file_size": file_path_obj.stat().st_size,
                "encoding": used_encoding,
            }

            # 计算文本哈希
            text_hash = hashlib.md5(text.encode()).hexdigest()
            metadata["text_hash"] = text_hash
            metadata["char_count"] = len(text)
            metadata["word_count"] = len(text.split())

            # 统计行数
            lines = text.split('\n')
            metadata["line_count"] = len(lines)

            logger.info(f"纯文本解析成功: {file_path_obj.name} (编码: {used_encoding}, 行数: {len(lines)})")

            return ParsedDocument(
                text=text,
                metadata=metadata,
                pages=[],
                tables=[],
                images=[]
            )

        except Exception as e:
            logger.error(f"纯文本解析失败: {file_path}, 错误: {str(e)}")
            raise


class HTMLParser(DocumentParser):
    """HTML文档解析器"""

    def __init__(self):
        try:
            from bs4 import BeautifulSoup
            self.BeautifulSoup = BeautifulSoup
            logger.info("HTML解析器初始化成功")
        except ImportError:
            logger.warning("beautifulsoup4未安装，HTML解析功能不可用")
            self.BeautifulSoup = None

    def supports(self, file_type: str) -> bool:
        return file_type.lower() in ["html", "htm"]

    def parse(self, file_path: str) -> ParsedDocument:
        """解析HTML文档"""
        if not self.BeautifulSoup:
            raise ImportError("请安装 beautifulsoup4: pip install beautifulsoup4")

        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            soup = self.BeautifulSoup(html_content, 'html.parser')

            # 移除脚本和样式标签
            for script in soup(["script", "style"]):
                script.decompose()

            # 提取文本
            text = soup.get_text(separator='\n', strip=True)

            # 提取元数据
            metadata = {
                "file_type": "html",
                "file_name": file_path_obj.name,
                "file_size": file_path_obj.stat().st_size,
            }

            # 提取标题
            title_tag = soup.find('title')
            if title_tag:
                metadata["title"] = title_tag.get_text().strip()

            # 提取meta标签
            meta_tags = soup.find_all('meta')
            for meta in meta_tags:
                name = meta.get('name')
                content = meta.get('content')
                if name and content:
                    metadata[f"meta_{name}"] = content

            # 提取链接
            links = [a.get('href') for a in soup.find_all('a') if a.get('href')]
            metadata["link_count"] = len(links)

            # 计算文本哈希
            text_hash = hashlib.md5(text.encode()).hexdigest()
            metadata["text_hash"] = text_hash
            metadata["char_count"] = len(text)
            metadata["word_count"] = len(text.split())

            logger.info(f"HTML解析成功: {file_path_obj.name} (链接: {len(links)})")

            return ParsedDocument(
                text=text,
                metadata=metadata,
                pages=[],
                tables=[],
                images=[]
            )

        except Exception as e:
            logger.error(f"HTML解析失败: {file_path}, 错误: {str(e)}")
            raise


# 注册的解析器
_PARSER_REGISTRY: List[DocumentParser] = [
    PDFParser(),
    DocxParser(),
    TxtParser(),
    HTMLParser(),
]


def get_parser(file_type: str) -> DocumentParser:
    """根据文件类型获取对应的解析器

    Args:
        file_type: 文件类型

    Returns:
        文档解析器

    Raises:
        ValueError: 不支持的文件类型
    """
    for parser in _PARSER_REGISTRY:
        if parser.supports(file_type):
            return parser

    raise ValueError(f"不支持的文件类型: {file_type}")


def register_parser(parser: DocumentParser):
    """注册新的文档解析器

    Args:
        parser: 文档解析器实例
    """
    _PARSER_REGISTRY.append(parser)
    logger.info(f"已注册文档解析器: {parser.__class__.__name__}")
