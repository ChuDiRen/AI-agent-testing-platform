"""
文档解析器
支持：PDF、Word、HTML、TXT
"""
import hashlib
from typing import List, Dict, Optional
from pathlib import Path

try:
    import PyPDF2
    from pdfplumber import page as pdfplumber_page
except ImportError:
    raise ImportError("请安装依赖：pip install PyPDF2 pdfplumber python-docx beautifulsoup4")


class DocumentParser:
    """文档解析器"""
    
    def __init__(self):
        self.parsers = {
            "pdf": self._parse_pdf,
            "docx": self._parse_docx,
            "txt": self._parse_txt,
            "html": self._parse_html,
        }
    
    def parse(self, file_path: str, file_type: str) -> dict:
        """
        解析文档
        
        Args:
            file_path: 文件路径
            file_type: 文档类型（pdf/docx/txt/html）
        
        Returns:
            包含 text、tables、metadata 的字典
        """
        parser = self.parsers.get(file_type)
        if parser:
            return parser(file_path)
        return {
            "text": "",
            "tables": [],
            "metadata": {}
        }
    
    def _parse_pdf(self, file_path: str) -> dict:
        """解析 PDF 文件"""
        try:
            content = ""
            tables = []
            
            with pdfplumber.open(file_path) as pdf:
                # 提取文本
                for i, page in enumerate(pdf.pages):
                    text += page.extract_text()
                
                # 提取表格
                for table in pdf.extract_tables():
                    tables.append({
                        "page": page.page_number,
                        "rows": [
                            [
                                [cell for cell in row for row in table.extract()]
                            ]
                        ]
                    })
            
            # 提取元数据
            metadata = {
                "author": pdf.info.author if pdf.info.author else "",
                "creator": pdf.info.creator if pdf.info.creator else "",
                "producer": pdf.info.producer if pdf.info.producer else "",
                "creation_date": pdf.info.creation_date,
                "modification_date": pdf.info.modification_date,
                "keywords": pdf.info.keywords if pdf.info.keywords else "",
                "page_count": len(pdf.pages),
            }
            
            return {
                "text": content,
                "tables": tables,
                "metadata": metadata
            }
        except Exception as e:
            raise ValueError(f"PDF解析失败: {str(e)}")
    
    def _parse_docx(self, file_path: str) -> dict:
        """解析 Word 文档"""
        try:
            import docx
            from docx import Document
            
            doc = Document(file_path)
            
            # 提取段落
            paragraphs = []
            for para in doc.paragraphs:
                paragraphs.append(para.text)
            
            # 提取表格
            tables = []
            for table in doc.tables:
                tables.append({
                    "rows": [[cell.text for cell in row.cells for row in table.rows]
                })
            
            # 提取元数据
            metadata = {
                "author": doc.core_properties.author,
                "title": doc.core_properties.title,
                "created": doc.core_properties.created,
                "modified": doc.core_properties.modified,
                "pages": len(doc.paragraphs)
            }
            
            return {
                "text": "\n\n".join(paragraphs),
                "tables": tables,
                "metadata": metadata
            }
        except Exception as e:
            raise ValueError(f"Word解析失败: {str(e)}")
    
    def _parse_txt(self, file_path: str) -> dict:
        """解析 TXT 文件"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            
            metadata = {
                "encoding": "utf-8",
                "line_count": len(text.splitlines()),
                "size_bytes": Path(file_path).stat().st_size
            }
            
            return {
                "text": text,
                "tables": [],
                "metadata": metadata
            }
        except Exception as e:
            raise ValueError(f"TXT解析失败: {str(e)}")
    
    def _parse_html(self, file_path: str) -> dict:
        """解析 HTML 文件"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, "r", encoding="utf-8") as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, "html.parser")
            
            # 提取纯文本
            text = soup.get_text(separator=" ", strip=True)
            text = " ".join(text.split())  # 去除多余空行
            
            # 提取表格
            tables = []
            for table in soup.find_all("table"):
                rows = []
                for row in table.find_all("tr"):
                    cells = []
                    for cell in row.find_all(["td", "th"]):
                        cells.append(cell.get_text(strip=True))
                    rows.append(cells)
                tables.append(rows)
            
            # 提取元数据
            metadata = {
                "title": soup.title.string or "",
                "charset": soup.meta.get("charset", ""),
                "link": soup.find("link", rel="canonical")["href"]) if soup.find("link", rel="canonical") else "",
                "page_count": len(soup.find_all("table")),
            }
            
            return {
                "text": text,
                "tables": tables,
                "metadata": metadata
            }
        except Exception as e:
            raise ValueError(f"HTML解析失败: {str(e)}")
    
    def compute_content_hash(self, content: str) -> str:
        """计算内容哈希（用于去重）"""
        return hashlib.md5(content).hexdigest()
    
    def split_text_into_chunks(
        self, 
        text: str, 
        chunk_size: int = 512, 
        overlap: int = 64
    ) -> List[Dict]:
        """
        将文本分割成块
        
        Args:
            text: 文本内容
            chunk_size: 块大小（字符数）
            overlap: 重叠大小（字符数）
        
        Returns:
            块列表
        """
        chunks = []
        
        # 按段落分割
        paragraphs = text.split("\n\n")
        
        current_chunk = ""
        chunk_index = 0
        
        for para in paragraphs:
            if not para.strip():
                continue
            
            # 检查是否需要分割当前块
            if len(current_chunk) + len(para) <= chunk_size:
                current_chunk += para + "\n\n"
            else:
                chunks.append({
                    "chunk_text": current_chunk.strip(),
                    "chunk_index": chunk_index
                })
                current_chunk = para + "\n\n"
                chunk_index += 1
        
        # 添加最后一个块
        if current_chunk.strip():
            chunks.append({
                "chunk_text": current_chunk.strip(),
                "chunk_index": chunk_index
            })
        
        return chunks


def parse_document(file_path: str) -> dict:
    """
    解析文档（智能识别类型）

    Args:
        file_path: 文件路径
        
    Returns:
        包含 text、tables、metadata 的字典
    """
    # 判断文件类型
    ext = file_path.split(".")[-1].lower()
    
    parser = DocumentParser()
    return parser.parse(file_path, ext)


def extract_text_from_file(file_path: str) -> str:
    """从文件中提取文本"""
    result = parse_document(file_path)
    return result["text"]
