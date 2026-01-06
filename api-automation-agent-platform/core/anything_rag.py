"""
AnythingChatRAG Integration - Enhanced Multi-Modal RAG Engine

This module provides enhanced RAG capabilities with multi-modal content support,
knowledge graph construction, and advanced retrieval strategies.
"""
from typing import Any, Dict, List, Optional, Union, Tuple
import json
import uuid
import re
import hashlib
from datetime import datetime
from pathlib import Path
import asyncio
from dataclasses import dataclass
from enum import Enum

# Multi-modal content processing
try:
    import fitz  # PyMuPDF for PDF
    import docx  # python-docx for Word
    import pandas as pd  # Excel processing
    from PIL import Image  # Image processing
    import pytesseract  # OCR for images
    MULTIMODAL_AVAILABLE = True
except ImportError:
    MULTIMODAL_AVAILABLE = False

# Vector storage and embeddings
import numpy as np
from sentence_transformers import SentenceTransformer


class RetrievalMode(Enum):
    """Retrieval modes for RAG"""
    LOCAL = "local"
    GLOBAL = "global"
    HYBRID = "hybrid"
    NAIVE = "naive"
    MIX = "mix"
    BYPASS = "bypass"


@dataclass
class Entity:
    """Knowledge graph entity"""
    id: str
    name: str
    type: str
    description: str
    properties: Dict[str, Any]
    embeddings: Optional[List[float]] = None


@dataclass
class Relationship:
    """Knowledge graph relationship"""
    id: str
    source_id: str
    target_id: str
    type: str
    description: str
    properties: Dict[str, Any]


@dataclass
class TextChunk:
    """Text chunk with metadata"""
    id: str
    content: str
    source: str
    chunk_type: str
    metadata: Dict[str, Any]
    embeddings: Optional[List[float]] = None


@dataclass
class RetrievalResult:
    """RAG retrieval result"""
    query: str
    mode: RetrievalMode
    entities: List[Entity]
    relationships: List[Relationship]
    chunks: List[TextChunk]
    citations: List[Dict[str, Any]]
    confidence: float
    processing_time: float


class AnythingChatRAGEngine:
    """
    Enhanced AnythingChatRAG Engine
    
    Features:
    - Multi-modal content processing (PDF, Word, Excel, Images)
    - Knowledge graph construction and querying
    - 6 retrieval modes
    - Entity and relationship extraction
    - Citation tracking
    """

    def __init__(
        self,
        persist_directory: str = "./data/anything_rag",
        embedding_model: str = "sentence-transformers/all-MiniLM-L6-v2",
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        enable_multimodal: bool = True
    ):
        """Initialize AnythingChatRAG engine"""
        self.persist_directory = Path(persist_directory)
        self.persist_directory.mkdir(parents=True, exist_ok=True)
        
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.enable_multimodal = enable_multimodal and MULTIMODAL_AVAILABLE
        
        # Initialize embedding model
        self.embedding_model = SentenceTransformer(embedding_model)
        self.embedding_dim = self.embedding_model.get_sentence_embedding_dimension()
        
        # Knowledge storage
        self.entities: Dict[str, Entity] = {}
        self.relationships: List[Relationship] = []
        self.chunks: List[TextChunk] = []
        
        # Vector indices
        self.entity_embeddings: Dict[str, List[float]] = {}
        self.chunk_embeddings: Dict[str, List[float]] = {}
        
        # Document registry
        self.documents: Dict[str, Dict[str, Any]] = {}
        
        # Statistics
        self.stats = {
            "total_documents": 0,
            "total_entities": 0,
            "total_relationships": 0,
            "total_chunks": 0,
            "last_updated": datetime.utcnow().isoformat()
        }

    async def add_document(
        self,
        file_path: Union[str, Path],
        document_type: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        Add document to RAG knowledge base
        
        Args:
            file_path: Path to document file
            document_type: Type of document (auto-detected if None)
            metadata: Additional metadata
        
        Returns:
            Document ID
        """
        file_path = Path(file_path)
        doc_id = str(uuid.uuid4())
        
        # Detect document type if not provided
        if document_type is None:
            document_type = self._detect_document_type(file_path)
        
        # Process document based on type
        if document_type.lower() == "pdf":
            content = await self._process_pdf(file_path)
        elif document_type.lower() in ["docx", "word"]:
            content = await self._process_word(file_path)
        elif document_type.lower() in ["xlsx", "excel"]:
            content = await self._process_excel(file_path)
        elif document_type.lower() in ["jpg", "jpeg", "png", "bmp", "tiff"]:
            content = await self._process_image(file_path)
        else:
            content = await self._process_text(file_path)
        
        # Create text chunks
        chunks = self._create_chunks(content, str(file_path), document_type)
        
        # Extract entities and relationships
        entities, relationships = await self._extract_knowledge(content, str(file_path))
        
        # Generate embeddings
        await self._generate_embeddings(chunks, entities)
        
        # Store in knowledge base
        for chunk in chunks:
            self.chunks.append(chunk)
        
        for entity in entities:
            self.entities[entity.id] = entity
        
        for relationship in relationships:
            self.relationships.append(relationship)
        
        # Register document
        self.documents[doc_id] = {
            "file_path": str(file_path),
            "document_type": document_type,
            "metadata": metadata or {},
            "chunk_count": len(chunks),
            "entity_count": len(entities),
            "relationship_count": len(relationships),
            "added_at": datetime.utcnow().isoformat()
        }
        
        # Update statistics
        self._update_stats()
        
        return doc_id

    async def query(
        self,
        query: str,
        mode: RetrievalMode = RetrievalMode.MIX,
        top_k: int = 10,
        chunk_top_k: int = 5,
        enable_rerank: bool = True,
        filters: Optional[Dict[str, Any]] = None
    ) -> RetrievalResult:
        """
        Query the RAG knowledge base
        
        Args:
            query: Query string
            mode: Retrieval mode
            top_k: Number of entities to return
            chunk_top_k: Number of text chunks to return
            enable_rerank: Enable result reranking
            filters: Query filters
        
        Returns:
            RetrievalResult with entities, relationships, and chunks
        """
        start_time = datetime.utcnow()
        
        # Generate query embedding
        query_embedding = self.embedding_model.encode(query, convert_to_tensor=True)
        
        # Execute query based on mode
        if mode == RetrievalMode.LOCAL:
            entities, relationships, chunks = await self._local_query(
                query_embedding, top_k, chunk_top_k, filters
            )
        elif mode == RetrievalMode.GLOBAL:
            entities, relationships, chunks = await self._global_query(
                query_embedding, top_k, chunk_top_k, filters
            )
        elif mode == RetrievalMode.HYBRID:
            entities, relationships, chunks = await self._hybrid_query(
                query_embedding, top_k, chunk_top_k, filters
            )
        elif mode == RetrievalMode.NAIVE:
            entities, relationships, chunks = await self._naive_query(
                query, top_k, chunk_top_k, filters
            )
        elif mode == RetrievalMode.BYPASS:
            entities, relationships, chunks = await self._bypass_query(
                query, top_k, chunk_top_k, filters
            )
        else:  # MIX mode (default)
            entities, relationships, chunks = await self._mix_query(
                query_embedding, query, top_k, chunk_top_k, filters
            )
        
        # Rerank results if enabled
        if enable_rerank:
            entities, relationships, chunks = await self._rerank_results(
                query, entities, relationships, chunks
            )
        
        # Generate citations
        citations = self._generate_citations(entities, relationships, chunks)
        
        # Calculate confidence
        confidence = self._calculate_confidence(entities, relationships, chunks)
        
        processing_time = (datetime.utcnow() - start_time).total_seconds()
        
        return RetrievalResult(
            query=query,
            mode=mode,
            entities=entities,
            relationships=relationships,
            chunks=chunks,
            citations=citations,
            confidence=confidence,
            processing_time=processing_time
        )

    async def _detect_document_type(self, file_path: Path) -> str:
        """Detect document type from file extension"""
        extension = file_path.suffix.lower()
        
        type_mapping = {
            ".pdf": "pdf",
            ".docx": "docx",
            ".doc": "docx",
            ".xlsx": "xlsx",
            ".xls": "xlsx",
            ".jpg": "jpg",
            ".jpeg": "jpg",
            ".png": "png",
            ".bmp": "bmp",
            ".tiff": "tiff",
            ".txt": "text",
            ".md": "markdown"
        }
        
        return type_mapping.get(extension, "text")

    async def _process_pdf(self, file_path: Path) -> str:
        """Process PDF document"""
        if not self.enable_multimodal:
            return f"PDF document: {file_path.name} (multimodal processing disabled)"
        
        try:
            doc = fitz.open(file_path)
            text_content = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text_content.append(page.get_text())
                
                # Extract images if multimodal is enabled
                if self.enable_multimodal:
                    image_list = page.get_images()
                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image["image"]
                            
                            # OCR processing
                            image = Image.open(io.BytesIO(image_bytes))
                            ocr_text = pytesseract.image_to_string(image)
                            if ocr_text.strip():
                                text_content.append(f"[Image {page_num+1}-{img_index+1}]: {ocr_text}")
                        except:
                            continue
            
            doc.close()
            return "\n".join(text_content)
            
        except Exception as e:
            return f"Error processing PDF: {str(e)}"

    async def _process_word(self, file_path: Path) -> str:
        """Process Word document"""
        if not self.enable_multimodal:
            return f"Word document: {file_path.name} (multimodal processing disabled)"
        
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                text_content.append("[Table]: " + "\n".join(table_text))
            
            return "\n".join(text_content)
            
        except Exception as e:
            return f"Error processing Word document: {str(e)}"

    async def _process_excel(self, file_path: Path) -> str:
        """Process Excel document"""
        if not self.enable_multimodal:
            return f"Excel document: {file_path.name} (multimodal processing disabled)"
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text_content = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert DataFrame to text
                sheet_text = f"[Sheet: {sheet_name}]\n"
                sheet_text += df.to_string(index=False)
                text_content.append(sheet_text)
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            return f"Error processing Excel document: {str(e)}"

    async def _process_image(self, file_path: Path) -> str:
        """Process image document with OCR"""
        if not self.enable_multimodal:
            return f"Image document: {file_path.name} (multimodal processing disabled)"
        
        try:
            image = Image.open(file_path)
            ocr_text = pytesseract.image_to_string(image)
            
            # Extract basic image metadata
            metadata = {
                "format": image.format,
                "size": image.size,
                "mode": image.mode
            }
            
            content = f"[Image: {file_path.name}]\n"
            content += f"Metadata: {json.dumps(metadata)}\n"
            content += f"OCR Text:\n{ocr_text}"
            
            return content
            
        except Exception as e:
            return f"Error processing image: {str(e)}"

    async def _process_text(self, file_path: Path) -> str:
        """Process text document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['gbk', 'gb2312', 'latin-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except:
                    continue
            return f"Could not read text file: {file_path}"

    def _create_chunks(self, content: str, source: str, doc_type: str) -> List[TextChunk]:
        """Create text chunks from content"""
        chunks = []
        
        # Split content into chunks
        if doc_type.lower() in ["xlsx", "excel"]:
            # For Excel, chunk by rows/sections
            lines = content.split('\n')
            current_chunk = []
            
            for line in lines:
                if line.startswith('[Sheet:') and current_chunk:
                    # Save previous chunk
                    chunk_content = '\n'.join(current_chunk)
                    chunk_id = str(uuid.uuid4())
                    chunks.append(TextChunk(
                        id=chunk_id,
                        content=chunk_content,
                        source=source,
                        chunk_type="table_section",
                        metadata={"document_type": doc_type}
                    ))
                    current_chunk = [line]
                else:
                    current_chunk.append(line)
                
                # Check chunk size
                if len('\n'.join(current_chunk)) > self.chunk_size:
                    chunk_content = '\n'.join(current_chunk[:-1])
                    chunk_id = str(uuid.uuid4())
                    chunks.append(TextChunk(
                        id=chunk_id,
                        content=chunk_content,
                        source=source,
                        chunk_type="table_section",
                        metadata={"document_type": doc_type}
                    ))
                    current_chunk = current_chunk[-1:]
            
            # Add remaining content
            if current_chunk:
                chunk_content = '\n'.join(current_chunk)
                chunk_id = str(uuid.uuid4())
                chunks.append(TextChunk(
                    id=chunk_id,
                    content=chunk_content,
                    source=source,
                    chunk_type="table_section",
                    metadata={"document_type": doc_type}
                ))
        else:
            # For other documents, use sliding window
            start = 0
            content_length = len(content)
            
            while start < content_length:
                end = start + self.chunk_size
                chunk_content = content[start:end]
                
                if chunk_content.strip():
                    chunk_id = str(uuid.uuid4())
                    chunks.append(TextChunk(
                        id=chunk_id,
                        content=chunk_content,
                        source=source,
                        chunk_type="text",
                        metadata={"document_type": doc_type}
                    ))
                
                start = end - self.chunk_overlap
        
        return chunks

    async def _extract_knowledge(self, content: str, source: str) -> Tuple[List[Entity], List[Relationship]]:
        """Extract entities and relationships from content"""
        entities = []
        relationships = []
        
        # API endpoint patterns
        api_patterns = [
            r'(GET|POST|PUT|DELETE|PATCH)\s+(/[^\s]+)',
            r'(POST|GET|PUT|DELETE|PATCH)\s+(/[^\s]+)',
            r'([A-Z]+)\s+(/[^\s]+)',
        ]
        
        # Parameter patterns
        param_patterns = [
            r'(\w+):\s*(string|number|boolean|array|object)',
            r'(\w+)\s*:\s*["\']([^"\']+)["\']',
        ]
        
        # Extract entities
        for pattern in api_patterns:
            matches = re.finditer(pattern, content, re.IGNORECASE)
            for match in matches:
                method = match.group(1)
                path = match.group(2)
                entity_id = str(uuid.uuid4())
                
                entity = Entity(
                    id=entity_id,
                    name=f"{method} {path}",
                    type="API_ENDPOINT",
                    description=f"API endpoint: {method} {path}",
                    properties={
                        "method": method,
                        "path": path,
                        "source": source
                    }
                )
                entities.append(entity)
        
        # Extract parameters
        for pattern in param_patterns:
            matches = re.finditer(pattern, content, re.IGNORECASE)
            for match in matches:
                param_name = match.group(1)
                param_type = match.group(2) if len(match.groups()) > 1 else "unknown"
                entity_id = str(uuid.uuid4())
                
                entity = Entity(
                    id=entity_id,
                    name=param_name,
                    type="PARAMETER",
                    description=f"Parameter: {param_name} ({param_type})",
                    properties={
                        "param_type": param_type,
                        "source": source
                    }
                )
                entities.append(entity)
        
        # Create relationships between entities
        api_entities = [e for e in entities if e.type == "API_ENDPOINT"]
        param_entities = [e for e in entities if e.type == "PARAMETER"]
        
        for api_entity in api_entities:
            for param_entity in param_entities:
                # Simple heuristic: if parameter is mentioned near API endpoint
                if param_entity.properties.get("source") == api_entity.properties.get("source"):
                    relationship_id = str(uuid.uuid4())
                    relationship = Relationship(
                        id=relationship_id,
                        source_id=api_entity.id,
                        target_id=param_entity.id,
                        type="HAS_PARAMETER",
                        description=f"{api_entity.name} has parameter {param_entity.name}",
                        properties={}
                    )
                    relationships.append(relationship)
        
        return entities, relationships

    async def _generate_embeddings(self, chunks: List[TextChunk], entities: List[Entity]) -> None:
        """Generate embeddings for chunks and entities"""
        # Generate chunk embeddings
        chunk_texts = [chunk.content for chunk in chunks]
        chunk_embeddings = self.embedding_model.encode(chunk_texts)
        
        for chunk, embedding in zip(chunks, chunk_embeddings):
            chunk.embeddings = embedding.tolist()
            self.chunk_embeddings[chunk.id] = embedding.tolist()
        
        # Generate entity embeddings
        entity_texts = [f"{entity.name} {entity.description}" for entity in entities]
        entity_embeddings = self.embedding_model.encode(entity_texts)
        
        for entity, embedding in zip(entities, entity_embeddings):
            entity.embeddings = embedding.tolist()
            self.entity_embeddings[entity.id] = embedding.tolist()

    async def _local_query(
        self, query_embedding, top_k: int, chunk_top_k: int, filters: Optional[Dict[str, Any]]
    ) -> Tuple[List[Entity], List[Relationship], List[TextChunk]]:
        """Local retrieval mode - entity and relationship focused"""
        # Find similar entities
        entity_scores = []
        for entity_id, entity_embedding in self.entity_embeddings.items():
            similarity = self._cosine_similarity(query_embedding, entity_embedding)
            entity_scores.append((similarity, entity_id))
        
        entity_scores.sort(reverse=True)
        top_entities = [self.entities[entity_id] for _, entity_id in entity_scores[:top_k]]
        
        # Find related relationships
        related_relationships = []
        entity_ids = {e.id for e in top_entities}
        for relationship in self.relationships:
            if relationship.source_id in entity_ids or relationship.target_id in entity_ids:
                related_relationships.append(relationship)
        
        # Find relevant chunks
        chunk_scores = []
        for chunk_id, chunk_embedding in self.chunk_embeddings.items():
            similarity = self._cosine_similarity(query_embedding, chunk_embedding)
            chunk_scores.append((similarity, chunk_id))
        
        chunk_scores.sort(reverse=True)
        top_chunks = [chunk for _, chunk_id in chunk_scores[:chunk_top_k] for chunk in self.chunks if chunk.id == chunk_id]
        
        return top_entities, related_relationships, top_chunks

    async def _global_query(
        self, query_embedding, top_k: int, chunk_top_k: int, filters: Optional[Dict[str, Any]]
    ) -> Tuple[List[Entity], List[Relationship], List[TextChunk]]:
        """Global retrieval mode - knowledge graph exploration"""
        # Start with all entities, explore relationships
        all_entities = list(self.entities.values())
        
        # Calculate similarity scores
        entity_scores = []
        for entity in all_entities:
            if entity.embeddings:
                similarity = self._cosine_similarity(query_embedding, entity.embeddings)
                entity_scores.append((similarity, entity))
        
        entity_scores.sort(reverse=True)
        top_entities = [entity for _, entity in entity_scores[:top_k]]
        
        # Explore relationships graph
        entity_ids = {e.id for e in top_entities}
        explored_ids = set(entity_ids)
        related_relationships = []
        
        # BFS exploration of relationships
        queue = list(entity_ids)
        while queue and len(related_relationships) < top_k * 2:
            current_id = queue.pop(0)
            
            for relationship in self.relationships:
                if relationship.source_id == current_id and relationship.target_id not in explored_ids:
                    related_relationships.append(relationship)
                    explored_ids.add(relationship.target_id)
                    queue.append(relationship.target_id)
                elif relationship.target_id == current_id and relationship.source_id not in explored_ids:
                    related_relationships.append(relationship)
                    explored_ids.add(relationship.source_id)
                    queue.append(relationship.source_id)
        
        # Get chunks for explored entities
        explored_entity_ids = {e.id for e in top_entities}
        explored_chunks = [chunk for chunk in self.chunks if any(
            entity_id in chunk.content for entity_id in explored_entity_ids
        )]
        
        return top_entities, related_relationships[:top_k], explored_chunks[:chunk_top_k]

    async def _hybrid_query(
        self, query_embedding, top_k: int, chunk_top_k: int, filters: Optional[Dict[str, Any]]
    ) -> Tuple[List[Entity], List[Relationship], List[TextChunk]]:
        """Hybrid retrieval mode - combine local and global"""
        # Get results from both modes
        local_entities, local_relationships, local_chunks = await self._local_query(
            query_embedding, top_k, chunk_top_k, filters
        )
        global_entities, global_relationships, global_chunks = await self._global_query(
            query_embedding, top_k, chunk_top_k, filters
        )
        
        # Combine and deduplicate
        all_entities = local_entities + global_entities
        entity_ids = {e.id for e in all_entities}
        combined_entities = [e for e in all_entities if e.id in entity_ids]
        
        all_relationships = local_relationships + global_relationships
        relationship_ids = {r.id for r in all_relationships}
        combined_relationships = [r for r in all_relationships if r.id in relationship_ids]
        
        all_chunks = local_chunks + global_chunks
        chunk_ids = {c.id for c in all_chunks}
        combined_chunks = [c for c in all_chunks if c.id in chunk_ids]
        
        return combined_entities[:top_k], combined_relationships[:top_k], combined_chunks[:chunk_top_k]

    async def _naive_query(
        self, query: str, top_k: int, chunk_top_k: int, filters: Optional[Dict[str, Any]]
    ) -> Tuple[List[Entity], List[Relationship], List[TextChunk]]:
        """Naive retrieval mode - simple text matching"""
        query_lower = query.lower()
        
        # Find entities with text matching
        matching_entities = []
        for entity in self.entities.values():
            if (query_lower in entity.name.lower() or 
                query_lower in entity.description.lower()):
                matching_entities.append(entity)
        
        # Find chunks with text matching
        matching_chunks = []
        for chunk in self.chunks:
            if query_lower in chunk.content.lower():
                matching_chunks.append(chunk)
        
        # Find relationships for matching entities
        entity_ids = {e.id for e in matching_entities}
        matching_relationships = [
            r for r in self.relationships
            if r.source_id in entity_ids or r.target_id in entity_ids
        ]
        
        return matching_entities[:top_k], matching_relationships[:top_k], matching_chunks[:chunk_top_k]

    async def _bypass_query(
        self, query: str, top_k: int, chunk_top_k: int, filters: Optional[Dict[str, Any]]
    ) -> Tuple[List[Entity], List[Relationship], List[TextChunk]]:
        """Bypass retrieval mode - direct query processing"""
        # Create a mock entity for the query itself
        query_entity = Entity(
            id=str(uuid.uuid4()),
            name=f"Query: {query}",
            type="QUERY",
            description=f"Direct query: {query}",
            properties={"query": query}
        )
        
        # Return recent chunks as context
        recent_chunks = sorted(self.chunks, key=lambda c: c.metadata.get("added_at", ""), reverse=True)[:chunk_top_k]
        
        return [query_entity], [], recent_chunks

    async def _mix_query(
        self, query_embedding, query: str, top_k: int, chunk_top_k: int, filters: Optional[Dict[str, Any]]
    ) -> Tuple[List[Entity], List[Relationship], List[TextChunk]]:
        """Mix retrieval mode - weighted combination of all modes"""
        # Get results from all modes
        local_entities, local_relationships, local_chunks = await self._local_query(
            query_embedding, top_k, chunk_top_k, filters
        )
        global_entities, global_relationships, global_chunks = await self._global_query(
            query_embedding, top_k, chunk_top_k, filters
        )
        naive_entities, naive_relationships, naive_chunks = await self._naive_query(
            query, top_k, chunk_top_k, filters
        )
        
        # Weighted scoring (mix of all approaches)
        entity_scores = {}
        
        # Score from local query (weight: 0.4)
        for i, entity in enumerate(local_entities):
            score = (top_k - i) / top_k * 0.4
            entity_scores[entity.id] = entity_scores.get(entity.id, 0) + score
        
        # Score from global query (weight: 0.3)
        for i, entity in enumerate(global_entities):
            score = (top_k - i) / top_k * 0.3
            entity_scores[entity.id] = entity_scores.get(entity.id, 0) + score
        
        # Score from naive query (weight: 0.3)
        for i, entity in enumerate(naive_entities):
            score = (top_k - i) / top_k * 0.3
            entity_scores[entity.id] = entity_scores.get(entity.id, 0) + score
        
        # Sort by combined score
        sorted_entities = sorted(
            [(score, entity_id) for entity_id, score in entity_scores.items()],
            reverse=True
        )
        
        final_entities = [self.entities[entity_id] for _, entity_id in sorted_entities[:top_k]]
        
        # Get relationships and chunks for top entities
        entity_ids = {e.id for e in final_entities}
        final_relationships = [
            r for r in self.relationships
            if r.source_id in entity_ids or r.target_id in entity_ids
        ][:top_k]
        
        final_chunks = [
            chunk for chunk in self.chunks
            if any(entity_id in chunk.content for entity_id in entity_ids)
        ][:chunk_top_k]
        
        return final_entities, final_relationships, final_chunks

    async def _rerank_results(
        self, query: str, entities: List[Entity], relationships: List[Relationship], chunks: List[TextChunk]
    ) -> Tuple[List[Entity], List[Relationship], List[TextChunk]]:
        """Rerank results based on query relevance"""
        # Simple reranking based on query term frequency
        query_terms = set(query.lower().split())
        
        # Rerank entities
        entity_scores = []
        for entity in entities:
            content = f"{entity.name} {entity.description}".lower()
            score = sum(1 for term in query_terms if term in content)
            entity_scores.append((score, entity))
        
        entity_scores.sort(reverse=True)
        reranked_entities = [entity for _, entity in entity_scores]
        
        # Rerank chunks
        chunk_scores = []
        for chunk in chunks:
            content = chunk.content.lower()
            score = sum(1 for term in query_terms if term in content)
            chunk_scores.append((score, chunk))
        
        chunk_scores.sort(reverse=True)
        reranked_chunks = [chunk for _, chunk in chunk_scores]
        
        # Relationships keep original order (based on entities)
        return reranked_entities, relationships, reranked_chunks

    def _generate_citations(
        self, entities: List[Entity], relationships: List[Relationship], chunks: List[TextChunk]
    ) -> List[Dict[str, Any]]:
        """Generate citations for results"""
        citations = []
        
        # Entity citations
        for entity in entities:
            source = entity.properties.get("source", "unknown")
            citations.append({
                "type": "entity",
                "id": entity.id,
                "name": entity.name,
                "source": source,
                "snippet": entity.description[:100] + "..." if len(entity.description) > 100 else entity.description
            })
        
        # Chunk citations
        for chunk in chunks:
            citations.append({
                "type": "chunk",
                "id": chunk.id,
                "source": chunk.source,
                "snippet": chunk.content[:150] + "..." if len(chunk.content) > 150 else chunk.content
            })
        
        return citations

    def _calculate_confidence(
        self, entities: List[Entity], relationships: List[Relationship], chunks: List[TextChunk]
    ) -> float:
        """Calculate confidence score for results"""
        # Simple confidence calculation based on result count and quality
        entity_score = min(len(entities) / 5, 1.0) * 0.4
        relationship_score = min(len(relationships) / 3, 1.0) * 0.3
        chunk_score = min(len(chunks) / 5, 1.0) * 0.3
        
        return entity_score + relationship_score + chunk_score

    def _cosine_similarity(self, vec1, vec2) -> float:
        """Calculate cosine similarity between vectors"""
        if isinstance(vec1, np.ndarray):
            vec1 = vec1.tolist()
        if isinstance(vec2, np.ndarray):
            vec2 = vec2.tolist()
        
        dot_product = sum(a * b for a, b in zip(vec1, vec2))
        magnitude1 = sum(a * a for a in vec1) ** 0.5
        magnitude2 = sum(b * b for b in vec2) ** 0.5
        
        if magnitude1 == 0 or magnitude2 == 0:
            return 0.0
        
        return dot_product / (magnitude1 * magnitude2)

    def _update_stats(self) -> None:
        """Update statistics"""
        self.stats.update({
            "total_documents": len(self.documents),
            "total_entities": len(self.entities),
            "total_relationships": len(self.relationships),
            "total_chunks": len(self.chunks),
            "last_updated": datetime.utcnow().isoformat()
        })

    def get_stats(self) -> Dict[str, Any]:
        """Get knowledge base statistics"""
        return self.stats.copy()

    def export_knowledge_graph(self) -> Dict[str, Any]:
        """Export knowledge graph for visualization"""
        return {
            "entities": [
                {
                    "id": entity.id,
                    "name": entity.name,
                    "type": entity.type,
                    "description": entity.description,
                    "properties": entity.properties
                }
                for entity in self.entities.values()
            ],
            "relationships": [
                {
                    "id": rel.id,
                    "source": rel.source_id,
                    "target": rel.target_id,
                    "type": rel.type,
                    "description": rel.description,
                    "properties": rel.properties
                }
                for rel in self.relationships
            ],
            "statistics": self.stats
        }


# Global instance for MCP server integration
anything_rag_instance: Optional[AnythingChatRAGEngine] = None


def get_anything_rag_instance(**kwargs) -> AnythingChatRAGEngine:
    """Get or create global AnythingChatRAG instance"""
    global anything_rag_instance
    if anything_rag_instance is None:
        anything_rag_instance = AnythingChatRAGEngine(**kwargs)
    return anything_rag_instance
