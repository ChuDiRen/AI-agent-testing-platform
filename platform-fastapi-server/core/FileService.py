import os
import logging
from typing import Optional
import asyncio

logger = logging.getLogger(__name__)


class FileService:
    """文件处理服务 - 支持TXT/Word/PDF文件解析"""
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """
        从TXT文件提取文本
        
        Args:
            file_path: 文件路径
            
        Returns:
            文本内容
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            with open(file_path, 'r', encoding='gbk') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading TXT file: {str(e)}")
            raise
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        从Word文件提取文本
        
        Args:
            file_path: 文件路径
            
        Returns:
            文本内容
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    text.append(" | ".join(row_data))
            
            return "\n".join(text)
        except ImportError:
            logger.error("python-docx not installed")
            raise
        except Exception as e:
            logger.error(f"Error reading DOCX file: {str(e)}")
            raise
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """
        从PDF文件提取文本
        
        Args:
            file_path: 文件路径
            
        Returns:
            文本内容
        """
        try:
            import PyPDF2
            
            text = []
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text.append(page.extract_text())
            
            return "\n".join(text)
        except ImportError:
            logger.error("PyPDF2 not installed")
            raise
        except Exception as e:
            logger.error(f"Error reading PDF file: {str(e)}")
            raise
    
    @staticmethod
    def extract_text_from_file(file_path: str) -> str:
        """
        根据文件类型自动选择合适的提取方法
        
        Args:
            file_path: 文件路径
            
        Returns:
            文本内容
            
        Raises:
            ValueError: 不支持的文件类型
        """
        # 获取文件扩展名
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext == '.txt':
            return FileService.extract_text_from_txt(file_path)
        elif ext in ['.docx', '.doc']:
            return FileService.extract_text_from_docx(file_path)
        elif ext == '.pdf':
            return FileService.extract_text_from_pdf(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    @staticmethod
    async def extract_text_from_file_async(file_path: str) -> str:
        """
        异步版本的文件提取
        
        Args:
            file_path: 文件路径
            
        Returns:
            文本内容
        """
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, FileService.extract_text_from_file, file_path)
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        清理文本（移除过多空白，标准化换行等）
        
        Args:
            text: 原始文本
            
        Returns:
            清理后的文本
        """
        # 移除多个连续换行
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            stripped = line.strip()
            if stripped:
                cleaned_lines.append(stripped)
        
        return '\n'.join(cleaned_lines)
    
    @staticmethod
    def truncate_text(text: str, max_length: int = 5000) -> str:
        """
        截断文本，限制长度
        
        Args:
            text: 文本
            max_length: 最大长度
            
        Returns:
            截断后的文本
        """
        if len(text) > max_length:
            return text[:max_length] + "... (文本已截断)"
        return text
