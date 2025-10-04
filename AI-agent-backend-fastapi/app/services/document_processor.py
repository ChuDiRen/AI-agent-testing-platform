# Copyright (c) 2025 左岚. All rights reserved.
"""文档处理服务 - 解析、分块、向量化"""
from typing import List, Dict, Any, Optional, Tuple
import os
import re
from pathlib import Path
import chardet

# 文档解析库
from pypdf import PdfReader
from docx import Document as DocxDocument
import markdown
from bs4 import BeautifulSoup

# 文本分块
from langchain_text_splitters import RecursiveCharacterTextSplitter


class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self):
        self.supported_types = {
            'pdf': self._parse_pdf,
            'docx': self._parse_docx,
            'doc': self._parse_docx,
            'txt': self._parse_txt,
            'md': self._parse_markdown,
            'html': self._parse_html,
            'htm': self._parse_html
        }
    
    def is_supported(self, file_type: str) -> bool:
        """检查文件类型是否支持"""
        return file_type.lower() in self.supported_types
    
    def parse_file(self, file_path: str, file_type: str) -> Tuple[str, Dict[str, Any]]:
        """解析文件"""
        file_type = file_type.lower().replace('.', '')
        
        if not self.is_supported(file_type):
            raise ValueError(f"不支持的文件类型: {file_type}")
        
        parser = self.supported_types[file_type]
        return parser(file_path)
    
    def _parse_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析PDF文件"""
        try:
            reader = PdfReader(file_path)
            text = ""
            metadata = {
                "page_count": len(reader.pages),
                "title": reader.metadata.get("/Title", "") if reader.metadata else "",
                "author": reader.metadata.get("/Author", "") if reader.metadata else ""
            }
            
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip(), metadata
        except Exception as e:
            raise Exception(f"PDF解析失败: {str(e)}")
    
    def _parse_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析Word文档"""
        try:
            doc = DocxDocument(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or ""
            }
            
            return text.strip(), metadata
        except Exception as e:
            raise Exception(f"Word文档解析失败: {str(e)}")
    
    def _parse_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析文本文件"""
        try:
            # 检测编码
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'utf-8'
            
            # 读取文本
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                text = f.read()
            
            metadata = {
                "encoding": encoding,
                "line_count": len(text.split('\n'))
            }
            
            return text.strip(), metadata
        except Exception as e:
            raise Exception(f"文本文件解析失败: {str(e)}")
    
    def _parse_markdown(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析Markdown文件"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                md_text = f.read()
            
            # 转换为HTML再提取纯文本
            html = markdown.markdown(md_text)
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text()
            
            # 提取标题
            titles = re.findall(r'^#+\s+(.+)$', md_text, re.MULTILINE)
            
            metadata = {
                "format": "markdown",
                "heading_count": len(titles),
                "first_heading": titles[0] if titles else ""
            }
            
            return text.strip(), metadata
        except Exception as e:
            raise Exception(f"Markdown文件解析失败: {str(e)}")
    
    def _parse_html(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析HTML文件"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html = f.read()
            
            soup = BeautifulSoup(html, 'html.parser')
            
            # 移除script和style标签
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text()
            
            # 清理多余空白
            lines = (line.strip() for line in text.splitlines())
            text = '\n'.join(line for line in lines if line)
            
            metadata = {
                "format": "html",
                "title": soup.title.string if soup.title else ""
            }
            
            return text.strip(), metadata
        except Exception as e:
            raise Exception(f"HTML文件解析失败: {str(e)}")
    
    def split_text(
        self,
        text: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        separators: Optional[List[str]] = None
    ) -> List[str]:
        """文本分块"""
        if separators is None:
            # 中文优化的分隔符
            separators = [
                "\n\n",  # 段落
                "\n",    # 行
                "。",    # 中文句号
                "！",    # 中文感叹号
                "？",    # 中文问号
                "；",    # 中文分号
                ".",     # 英文句号
                "!",     # 英文感叹号
                "?",     # 英文问号
                " ",     # 空格
                ""       # 字符
            ]
        
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=separators,
            length_function=len,
            is_separator_regex=False
        )
        
        chunks = splitter.split_text(text)
        return chunks
    
    def process_document(
        self,
        file_path: str,
        file_type: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50
    ) -> Tuple[List[str], Dict[str, Any]]:
        """处理文档: 解析 + 分块"""
        # 解析文档
        text, metadata = self.parse_file(file_path, file_type)
        
        # 分块
        chunks = self.split_text(text, chunk_size, chunk_overlap)
        
        # 更新元数据
        metadata.update({
            "total_chars": len(text),
            "chunk_count": len(chunks),
            "chunk_size": chunk_size,
            "chunk_overlap": chunk_overlap
        })
        
        return chunks, metadata
    
    def extract_keywords(self, text: str, top_k: int = 10) -> List[str]:
        """提取关键词(简单实现)"""
        # 移除标点符号
        text = re.sub(r'[^\w\s]', ' ', text)
        
        # 分词(简单按空格分)
        words = text.split()
        
        # 统计词频
        word_freq = {}
        for word in words:
            word = word.lower()
            if len(word) > 1:  # 过滤单字符
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # 排序
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        
        return [word for word, freq in sorted_words[:top_k]]
    
    def calculate_stats(self, text: str) -> Dict[str, Any]:
        """计算文本统计信息"""
        return {
            "char_count": len(text),
            "word_count": len(text.split()),
            "line_count": len(text.split('\n')),
            "paragraph_count": len([p for p in text.split('\n\n') if p.strip()])
        }


# 全局文档处理器实例
document_processor = DocumentProcessor()

